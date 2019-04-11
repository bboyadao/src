import shopify
import ssl
from docx import Document
from docxcompose.composer import Composer

ssl._create_default_https_context = ssl._create_unverified_context
API_KEY, PASSWORD = "6413e468090c0e3acac67e564db39752", "7bd6e0bffac90790b6a6343d3042e923"
shop_url = "https://%s:%s@kourse-com.myshopify.com/admin" % (
    API_KEY, PASSWORD)

shopify.ShopifyResource.set_site(shop_url)


def get_all_resources(resource, **kwargs):
    resource_count = resource.count(**kwargs)
    resources = []
    if resource_count > 0:
        for page in range(1, ((resource_count - 1) // 250) + 2):
            kwargs.update({"limit": 250, "page": page})
            resources.extend(resource.find(**kwargs))
    return resources


def list_products():

    shop = shopify.Shop.current
    collec = shopify.CollectionListing
    products = get_all_resources(shopify.Product)

    return products


def create_collection(resources, cm_name, logo):
    collect = shopify.Collect(
        {'product_id': product_id, 'collection_id': collection_id})
    collect.save()
    pass


def create_product():
    new_product = shopify.Product()
    new_product.title = "Burton Custom Freestyle 151"
    new_product.product_type = "Snowboard"
    new_product.body_html = "<strong>Good snowboard!</strong>"
    new_product.vendor = "Burton"
    image1 = shopify.Image()
    image1.src = logo

    new_product.images = [image1]
    new_product.save()
    success = new_product.save()


def merge():
    files = [
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 1.docx',
        "/Users/xiaoma/projects/shopify-book/src/Templates/Study Guide - Template 1 - tadaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa.docx",
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/amazon.docx', '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 2.docx',
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Internet Commerce.docx',
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 3.docx',
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Amazon.docx',
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 4.docx',
        # '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Amazon Fulfillment Associate.docx',
        '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 5.docx',
        #  '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/amazon_operations.docx', '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 6.docx'
    ]
    merged_name = "/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/tabbbbbbbbb.docx"
    master = Document(files[0])
    master1 = Document(files[1])
    composer = Composer(master)
    composer.append(master1)

    composer.save(merged_name)

    # composer.append(doc1, doc2)


def mer2():
    temp1 = '/Users/xiaoma/projects/shopify-book/src/Templates/Study Guide - Template 1.docx'
    doc = Document(temp1)
    return doc
    foot = doc.sections[-1].footer

    merged_name = "/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/tabbbbbbbbb.docx"
    # doc1 = Document(merged_name).sections[-1].footer = foot
    # doc1
    # return doc
    # print(doc.__dict__)


def mer3():
    temp1 = '/Users/xiaoma/projects/shopify-book/src/Templates/Study Guide - Template 1.docx'
    doc = Document(temp1)
    a = doc._body._body
    print(dir(a))
    for i in a:
        print(i.text)

    # a.replace("[", "fuckkkkkkkkk")
    doc.save('/Users/xiaoma/projects/shopify-book/src/Templates/Study Guide - Template 1 - tadaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa.docx')


if __name__ == '__main__':
    # a = merge()

    # merge()
    merged_name = '/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/Study Guide - Template 1.docx'
    docx = mer2()
    a = docx.sections[-1]
    foot_src = a.footer
    doc = Document(merged_name)
    head = doc.sections[-1].header
    head.add_header("ssssss")
    footer = doc.sections[-1].footer
    footer.paragraphs[0].text = "fuckkkkkkkkk"
    # footer.add_paragraph(
    # "ffffffffffffffffffffffffffffffffffffffffffffffffffffff")
    doc.add_heading('Document Title', 0)

    doc.save("/Users/xiaoma/projects/shopify-book/src/Output/Amazon/Fulfillment Associate/Temp Files/changefottt.docx")
